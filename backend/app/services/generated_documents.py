from io import BytesIO

from docx import Document as DocxDocument

from app.storage.local import local_storage


async def save_generated_document_text(content: str) -> str:
    return await local_storage.save_text(content, suffix=".txt")


async def build_docx_export(title: str, content: str) -> tuple[bytes, str]:
    document = DocxDocument()
    document.add_heading(title, level=1)
    for block in content.split("\n"):
        if block.strip():
            document.add_paragraph(block)
        else:
            document.add_paragraph("")
    buffer = BytesIO()
    document.save(buffer)
    data = buffer.getvalue()
    storage_key = await local_storage.save(data, suffix=".docx")
    return data, storage_key


async def build_pdf_export(title: str, content: str) -> tuple[bytes, str]:
    data = _minimal_pdf_bytes(title, content)
    storage_key = await local_storage.save(data, suffix=".pdf")
    return data, storage_key


def _minimal_pdf_bytes(title: str, content: str) -> bytes:
    lines = [title, "", *content.split("\n")]
    y = 800
    commands = ["BT", "/F1 12 Tf"]
    for line in lines[:45]:
        safe_line = _escape_pdf_text(line[:110])
        commands.append(f"72 {y} Td ({safe_line}) Tj")
        commands.append(f"-72 -16 Td")
        y -= 16
    commands.append("ET")
    stream = "\n".join(commands).encode("utf-8")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        b"5 0 obj << /Length " + str(len(stream)).encode("ascii") + b" >> stream\n" + stream + b"\nendstream endobj\n",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return bytes(pdf)


def _escape_pdf_text(value: str) -> str:
    ascii_value = value.encode("ascii", "replace").decode("ascii")
    return ascii_value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
