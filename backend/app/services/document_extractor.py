import base64
import io
from dataclasses import dataclass
from pathlib import Path

import httpx

from app.core.config import settings


IMAGE_MIME_TYPES = {"image/jpeg", "image/png", "image/webp"}


class DocumentExtractionError(RuntimeError):
    pass


@dataclass
class ExtractionResult:
    text: str
    extraction_status: str
    ocr_status: str
    error: str | None = None


class VisionOCRClient:
    async def extract_text(self, content: bytes, mime_type: str) -> str:
        if not settings.document_vision_model or not settings.document_vision_provider:
            raise DocumentExtractionError("Vision OCR model is not configured.")
        if not settings.openrouter_api_key:
            raise DocumentExtractionError("OPENROUTER_API_KEY is not configured for vision OCR.")

        data_url = f"data:{mime_type};base64,{base64.b64encode(content).decode('ascii')}"
        payload = {
            "model": settings.document_vision_model,
            "provider": {"only": [settings.document_vision_provider], "require_parameters": True},
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract all readable text from this business document image. Return text only."},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }
            ],
            "max_tokens": min(settings.openrouter_max_output_tokens, 1200),
        }
        headers = {
            "Authorization": f"Bearer {settings.openrouter_api_key}",
            "Content-Type": "application/json",
            "X-OpenRouter-Title": settings.openrouter_app_title,
        }
        async with httpx.AsyncClient(timeout=settings.openrouter_timeout_seconds) as client:
            response = await client.post(f"{settings.openrouter_base_url}/chat/completions", headers=headers, json=payload)
            response.raise_for_status()
        data = response.json()
        return ((data.get("choices") or [{}])[0].get("message") or {}).get("content") or ""


class DocumentExtractor:
    def __init__(self, vision_client: VisionOCRClient | None = None) -> None:
        self.vision_client = vision_client or VisionOCRClient()

    async def extract(self, content: bytes, mime_type: str, filename: str) -> ExtractionResult:
        try:
            if mime_type == "text/plain":
                return ExtractionResult(text=self._extract_txt(content), extraction_status="completed", ocr_status="not_required")
            if mime_type == "application/pdf":
                text = self._extract_pdf(content)
                if not text.strip():
                    return ExtractionResult(text="", extraction_status="failed", ocr_status="pending", error="PDF has no extractable text.")
                return ExtractionResult(text=text, extraction_status="completed", ocr_status="not_required")
            if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return ExtractionResult(text=self._extract_docx(content), extraction_status="completed", ocr_status="not_required")
            if mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                return ExtractionResult(text=self._extract_xlsx(content), extraction_status="completed", ocr_status="not_required")
            if mime_type in IMAGE_MIME_TYPES:
                text = await self.vision_client.extract_text(content, mime_type)
                return ExtractionResult(text=text, extraction_status="completed", ocr_status="completed")
        except UnicodeDecodeError as exc:
            return ExtractionResult(text="", extraction_status="failed", ocr_status="not_required", error="TXT file must be UTF-8.")
        except Exception as exc:
            ocr_status = "failed" if mime_type in IMAGE_MIME_TYPES else "not_required"
            return ExtractionResult(text="", extraction_status="failed", ocr_status=ocr_status, error=str(exc))

        raise DocumentExtractionError(f"Unsupported document type for {Path(filename).suffix}")

    def _extract_txt(self, content: bytes) -> str:
        return content.decode("utf-8")

    def _extract_pdf(self, content: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join(page for page in pages if page)

    def _extract_docx(self, content: bytes) -> str:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    lines.append(" | ".join(values))
        return "\n".join(lines)

    def _extract_xlsx(self, content: bytes) -> str:
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        lines: list[str] = []
        max_rows = settings.xlsx_max_rows
        for sheet in workbook.worksheets:
            lines.append(f"[Sheet: {sheet.title}]")
            for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                if index > max_rows:
                    lines.append(f"[Rows truncated at {max_rows}]")
                    break
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    lines.append("\t".join(values))
        return "\n".join(lines)


document_extractor = DocumentExtractor()
